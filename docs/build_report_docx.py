#!/usr/bin/env python3
"""
Build docs/Capstone_Report.docx from docs/Capstone_Report.md  (PREVIEW build).

Preview intent:
  - Real heading styles (Title / Subtitle / Heading 1 / Heading 2) so the ToC + outline work.
  - A live, updatable Table of Contents field (Word: right-click -> Update Field).
  - Page numbers in the footer ("Page X of Y").
  - All 12 figures embedded (../outputs/*.png and ../data/processed/*.png resolved from docs/).
  - Pipe tables -> real Word tables (gridlines, shaded bold header, column alignment).
  - Inline **bold** / *italic* preserved.
  - The ~66 [CITE: ...], [VERIFY: ...], [TEMPLATE] markers are LEFT VISIBLE as literal text.
  - The top HTML drafting-conventions comment block is skipped.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

MD_PATH = "/Users/feya/Downloads/My capstone project/docs/Capstone_Report.md"
OUT_PATH = "/Users/feya/Downloads/My capstone project/docs/Capstone_Report.docx"
DOCS_DIR = os.path.dirname(MD_PATH)

INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*([^*]+?)\*")
IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
META_LABELS = ("**Author:**", "**Programme:**", "**Capstone client:**", "**Date:**")

stats = {"headings": 0, "h1": 0, "h2": 0, "images": 0, "tables": 0,
         "missing_images": 0, "paragraphs": 0, "bullets": 0, "captions": 0}


def add_inline(paragraph, text, bold=False, italic=False, size=None):
    """Render a markdown string into runs, honouring **bold** and *italic*."""
    def style(run, b, i):
        if b:
            run.bold = True
        if i:
            run.italic = True
        if size:
            run.font.size = size

    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            style(paragraph.add_run(text[pos:m.start()]), bold, italic)
        if m.group(1) is not None:        # **bold**
            style(paragraph.add_run(m.group(1)), True, italic)
        else:                              # *italic*
            style(paragraph.add_run(m.group(2)), bold, True)
        pos = m.end()
    if pos < len(text):
        style(paragraph.add_run(text[pos:]), bold, italic)


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_sep(line):
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = split_row(s)
    return len(cells) > 0 and all(re.match(r"^:?-+:?$", c) for c in cells)


def align_of(sepcell):
    sc = sepcell.strip()
    left = sc.startswith(":")
    right = sc.endswith(":")
    if left and right:
        return WD_ALIGN_PARAGRAPH.CENTER
    if right:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, rows):
    header = split_row(rows[0])
    aligns = [align_of(c) for c in split_row(rows[1])]
    data = [split_row(r) for r in rows[2:]]
    ncols = len(header)
    t = doc.add_table(rows=1 + len(data), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    for j, ctext in enumerate(header):
        cell = t.cell(0, j)
        p = cell.paragraphs[0]
        p.alignment = aligns[j] if j < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, ctext, bold=True)
        shade_cell(cell, "D9E2F3")

    for r_i, cells in enumerate(data):
        for j in range(ncols):
            cell = t.cell(r_i + 1, j)
            p = cell.paragraphs[0]
            p.alignment = aligns[j] if j < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
            if j < len(cells):
                add_inline(p, cells[j])
    stats["tables"] += 1
    # spacer paragraph after a table
    doc.add_paragraph()


def add_image(doc, rel_path, alt):
    abs_path = os.path.normpath(os.path.join(DOCS_DIR, rel_path))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not os.path.exists(abs_path):
        add_inline(p, f"[Figure not found: {rel_path}]", italic=True)
        stats["missing_images"] += 1
        return
    try:
        with Image.open(abs_path) as img:
            w, h = img.size
        aspect = h / w
        max_w, max_h = 6.3, 7.8
        run = p.add_run()
        if max_w * aspect <= max_h:
            run.add_picture(abs_path, width=Inches(max_w))
        else:
            run.add_picture(abs_path, height=Inches(max_h))
        stats["images"] += 1
    except Exception as e:  # noqa
        add_inline(p, f"[Figure could not be embedded: {rel_path} ({e})]", italic=True)
        stats["missing_images"] += 1


def add_caption(doc, text):
    inner = text[1:-1] if text.startswith("*") and text.endswith("*") else text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, inner, italic=True, size=Pt(9))
    stats["captions"] += 1


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    placeholder = p.add_run(
        "Right-click and choose “Update Field” to build the table of "
        "contents (page numbers appear after updating)."
    )
    placeholder.italic = True
    run_end = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_page_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")

    def field(instr_text):
        r = p.add_run()
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr_text
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        r._r.append(b); r._r.append(it); r._r.append(e)

    field("PAGE")
    p.add_run(" of ")
    field("NUMPAGES")


def build():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    # US Letter page setup + 1in margins + footer page numbers
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(1))
    add_page_footer(section)

    title_seen = False
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        s = raw.strip()

        # --- skip HTML comment block ---
        if s.startswith("<!--"):
            while i < n and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if s == "":
            i += 1
            continue

        # horizontal rules: chapter breaks are handled via page_break_before
        if s == "---":
            i += 1
            continue

        # headings
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            text = s[level:].strip()
            if level == 1 and not title_seen:
                h = doc.add_heading(text, 0)            # Title style
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_seen = True
            elif level == 1:
                h = doc.add_heading(text, 1)
                h.paragraph_format.page_break_before = True
                stats["h1"] += 1
                stats["headings"] += 1
            elif level == 2:
                doc.add_heading(text, 2)
                stats["h2"] += 1
                stats["headings"] += 1
            elif level == 3:
                p = doc.add_paragraph(text, style="Subtitle")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, min(level, 4))
                stats["headings"] += 1
            i += 1
            continue

        # ToC placeholder line -> live TOC field
        if s.startswith("[Auto-generated at typeset"):
            add_toc(doc)
            i += 1
            continue

        # standalone image
        m = IMG_RE.match(s)
        if m:
            add_image(doc, m.group(2), m.group(1))
            i += 1
            continue

        # figure caption (centered, small, italic)
        if s.startswith("*Figure ") and s.endswith("*"):
            add_caption(doc, s)
            i += 1
            continue

        # table block
        if s.startswith("|") and (i + 1 < n) and is_sep(lines[i + 1]):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            continue

        # bullet list item
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, s[2:].strip())
            stats["bullets"] += 1
            i += 1
            continue

        # centered title-page metadata
        if any(s.startswith(lbl) for lbl in META_LABELS):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, s)
            i += 1
            continue

        # normal body paragraph
        p = doc.add_paragraph()
        # whole-line italic notes / bold captions read fine; justify body prose
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, s)
        stats["paragraphs"] += 1
        i += 1

    # python-docx's bundled default template ships <w:zoom/> without the
    # schema-required w:percent attribute -> add it so the file validates.
    settings_el = doc.settings.element
    zoom = settings_el.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    doc.save(OUT_PATH)
    return doc


if __name__ == "__main__":
    build()
    # Re-open to validate structural integrity
    d2 = Document(OUT_PATH)
    size = os.path.getsize(OUT_PATH)
    print("=== BUILD COMPLETE ===")
    print(f"output      : {OUT_PATH}")
    print(f"size        : {size:,} bytes")
    print(f"paragraphs  : {len(d2.paragraphs)}")
    print(f"inline_imgs : {len(d2.inline_shapes)}")
    print(f"tables      : {len(d2.tables)}")
    print("--- counters ---")
    for k, v in stats.items():
        print(f"{k:14}: {v}")
